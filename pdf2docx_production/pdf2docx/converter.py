"""Core PDF to DOCX conversion functionality."""

import sys
import subprocess
import os
import tempfile
import shutil
import logging
import difflib
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Tuple, Dict, Any, Optional

try:
    from IPython import get_ipython
    IPYTHON_AVAILABLE = True
except ImportError:
    IPYTHON_AVAILABLE = False

# PDF processing libraries
REQUIRED = {"pdf2docx": "pdf2docx", "docx": "python-docx", "PIL": "Pillow"}
OPTIONAL = {"ocrmypdf": "ocrmypdf"}

def _pip_install(pip_names: List[str]) -> None:
    """Install required Python packages with multiple fallback methods."""
    attempts = []
    probe = subprocess.run([sys.executable, "-m", "pip", "--version"], capture_output=True)
    if probe.returncode != 0:
        subprocess.run([sys.executable, "-m", "ensurepip", "--upgrade"], capture_output=True)

    proc = subprocess.run(
        [sys.executable, "-m", "pip", "install", "--break-system-packages", *pip_names],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    if proc.returncode == 0:
        return
    attempts.append(("subprocess: python -m pip", (proc.stdout + proc.stderr).strip()))

    if IPYTHON_AVAILABLE:
        try:
            from IPython.utils.capture import capture_output
            ip = get_ipython()
            if ip is not None:
                with capture_output() as cap:
                    ip.run_line_magic("pip", "install --break-system-packages " + " ".join(pip_names))
                text = (cap.stdout or "") + (cap.stderr or "")
                if "successfully installed" in text.lower() or "requirement already satisfied" in text.lower():
                    return
                attempts.append(("%pip install (notebook kernel)", text.strip()))
        except Exception as e:
            attempts.append(("%pip install (notebook kernel)", f"could not run: {e}"))

    report = "\n\n".join(f"--- Attempt: {name} ---\n" + "\n".join(out.splitlines()[-12:])
                          for name, out in attempts)
    raise RuntimeError(
        f"Could not install {pip_names} — tried {len(attempts)} method(s), all failed.\n\n"
        f"{report}\n\n"
        "You can also try installing manually in a terminal:\n"
        f"  {sys.executable} -m pip install {' '.join(pip_names)}"
    )

def _cannot_import(mod: str) -> bool:
    """Check if a module can be imported."""
    try:
        __import__(mod)
        return False
    except ImportError:
        return True

def _ensure_dependencies(verbose: bool = True) -> Dict[str, bool]:
    """Ensure all required dependencies are installed and check optional ones."""
    missing_required = [pip for mod, pip in REQUIRED.items() if _cannot_import(mod)]
    if missing_required:
        if verbose:
            print(f"Installing required packages: {missing_required} ...")
        _pip_install(missing_required)
        still_missing = [mod for mod in REQUIRED if _cannot_import(mod)]
        if still_missing:
            raise ImportError(
                f"{still_missing} installed without error but still aren't importable — "
                "this usually means installation landed in a different Python environment "
                f"than the one running this notebook ({sys.executable}). Try restarting "
                "the kernel and running this cell again."
            )

    ocr_available = not _cannot_import("ocrmypdf")
    if not ocr_available:
        try:
            _pip_install([OPTIONAL["ocrmypdf"]])
            ocr_available = not _cannot_import("ocrmypdf")
        except Exception:
            ocr_available = False

    caps = {
        "ocrmypdf_installed": ocr_available,
        "tesseract": shutil.which("tesserat") is not None,
        "ghostscript": shutil.which("gs") is not None,
        "unpaper": shutil.which("unpaper") is not None,
        "libreoffice": shutil.which("soffice") is not None,
    }
    return caps

# Initialize dependencies and logger
CAPS = _ensure_dependencies(verbose=False)

try:
    import pymupdf as fitz
except ImportError as e:
    if _cannot_import("pymupdf"):
        _pip_install(["pymupdf>=1.24"])
        import pymupdf as fitz
    else:
        raise

# Set up fitz module alias for backward compatibility
sys.modules.setdefault("fitz", fitz)

try:
    import ocrmypdf
    from ocrmypdf.exceptions import MissingDependencyError as _OcrMissingDep
except ImportError:
    ocrmypdf = None
    _OcrMissingDep = Exception

from docx import Document as DocxDocument
from PIL import Image

# Configure logging
logger = logging.getLogger("pdf2docx_converter")
if not logger.handlers:  # Avoid adding handlers multiple times
    handler = logging.StreamHandler()
    formatter = logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")
    handler.setFormatter(formatter)
    logger.addHandler(handler)
    logger.setLevel(logging.INFO)

# Constants
SPARSE_CHAR_THRESHOLD = 40

# Converter will be imported lazily to avoid circular imports
_Converter = None

def _get_Converter():
    """Get the Converter class from the external pdf2docx package."""
    global _Converter
    if _Converter is None:
        try:
            # Import the external pdf2docx package directly via file path to avoid
            # conflicts with our own package which has the same name
            import importlib.util
            import sys

            # Path to the external pdf2docx package
            external_pdf2docx_path = r"C:\Python314\Lib\site-packages\pdf2docx\__init__.py"

            # Load the module
            spec = importlib.util.spec_from_file_location("pdf2docx_external", external_pdf2docx_path)
            _external_pdf2docx = importlib.util.module_from_spec(spec)

            # Add to sys.modules to prevent reload issues
            sys.modules["pdf2docx_external"] = _external_pdf2docx

            # Execute the module
            spec.loader.exec_module(_external_pdf2docx)

            # Get the Converter class
            _Converter = _external_pdf2docx.Converter
            print(f"[DEBUG] Successfully loaded external pdf2docx Converter: {_Converter}")
        except Exception as e:
            print(f"[DEBUG] Error loading external pdf2docx: {e}")
            # Fallback to try normal import (in case the external package is installed differently)
            try:
                import pdf2docx as _external_pdf2docx
                _Converter = _external_pdf2docx.Converter
                print(f"[DEBUG] Fallback import successful: {_Converter}")
            except Exception as e2:
                print(f"[DEBUG] Fallback import also failed: {e2}")
                raise e  # Raise the original exception
    return _Converter

def _open_pdf(pdf_path: str, password: Optional[str] = None) -> fitz.Document:
    """Open a PDF file with optional password protection."""
    doc = fitz.open(pdf_path)
    if doc.needs_pass and not doc.authenticate(password or ""):
        doc.close()
        raise ValueError(f"'{pdf_path}' is password-protected and the password given was incorrect (or none was given).")
    return doc

def page_text_density(pdf_path: str, password: Optional[str] = None) -> List[int]:
    """Get text density (character count) for each page in a PDF."""
    doc = _open_pdf(pdf_path, password)
    try:
        return [len(page.get_text()) for page in doc]
    finally:
        doc.close()

def sparse_page_indices(pdf_path: str, chars_per_page_threshold: int = SPARSE_CHAR_THRESHOLD,
                         password: Optional[str] = None) -> List[int]:
    """Find indices of pages with text density below threshold (likely scanned/image-only)."""
    return [i for i, c in enumerate(page_text_density(pdf_path, password)) if c < chars_per_page_threshold]

def is_scanned(pdf_path: str, chars_per_page_threshold: int = SPARSE_CHAR_THRESHOLD,
                fraction_threshold: float = 0.6, password: Optional[str] = None) -> bool:
    """Determine if a PDF is primarily scanned (image-based) rather than text-based."""
    try:
        counts = page_text_density(pdf_path, password)
        if not counts:
            return False
        sparse = sum(1 for c in counts if c < chars_per_page_threshold)
        return (sparse / len(counts)) >= fraction_threshold
    except Exception as e:
        logger.warning(f"Could not analyze {pdf_path}: {e}")
        return False

def ocr_pdf(input_pdf: str, output_pdf: str, *, language: str = "eng", deskew: bool = True,
            clean: bool = False, force_ocr: bool = False, **ocr_kwargs) -> bool:
    """Apply OCR to a PDF file using ocrmypdf."""
    if ocrmypdf is None or not CAPS["tesseract"]:
        logger.error("OCR unavailable: install 'ocrmypdf' (pip) and the Tesseract engine (system package).")
        return False

    clean = clean and CAPS["unpaper"]
    if ocr_kwargs.get("clean_final") and not CAPS["unpaper"]:
        ocr_kwargs.pop("clean_final", None)

    try:
        logger.info(f"Running OCR on: {input_pdf}")
        ocrmypdf.ocr(input_pdf, output_pdf, language=language, deskew=deskew, clean=clean,
                     skip_text=not force_ocr, force_ocr=force_ocr, progress_bar=False, **ocr_kwargs)
        logger.info("OCR completed successfully.")
        return True
    except _OcrMissingDep as e:
        logger.error(f"OCR failed — missing system dependency: {e}")
        return False
    except Exception as e:
        logger.error(f"OCR failed: {e}")
        return False

@dataclass
class FidelityReport:
    """Report on the fidelity of PDF to DOCX conversion."""
    pdf_pages: int
    docx_rendered_pages: Optional[int]
    text_similarity: Optional[float]
    pdf_char_count: int
    docx_char_count: int
    page_images: List[Tuple[str, str]] = field(default_factory=list)
    comparison_images: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)

    def summary(self) -> str:
        """Generate a human-readable summary of the fidelity report."""
        sim = f"{self.text_similarity*100:.1f}%" if self.text_similarity is not None else "n/a"
        lines = [
            f"Text similarity: {sim}",
            f"Pages — PDF: {self.pdf_pages}, DOCX (rendered): {self.docx_rendered_pages if self.docx_rendered_pages is not None else 'n/a'}",
            f"Characters — PDF: {self.pdf_char_count}, DOCX: {self.docx_char_count}"
        ]
        for w in self.warnings:
            lines.append(f"Note: {w}")
        return "\n".join(lines)

@dataclass
class ConversionResult:
    """Result of a PDF to DOCX conversion attempt."""
    pdf_path: str
    docx_path: Optional[str]
    success: bool
    used_ocr: bool = False
    ocr_supplement_pages: int = 0
    error: Optional[str] = None
    fidelity: Optional[FidelityReport] = None

    def __post_init__(self):
        # Set success based on whether we have a docx_path and no error
        if self.success is None:  # If not explicitly set
            self.success = self.docx_path is not None and self.error is None

    def summary(self) -> str:
        """Generate a human-readable summary of the conversion result."""
        if not self.success:
            return f"FAILED: {self.pdf_path} -> {self.error}"
        s = f"OK: {self.pdf_path} -> {self.docx_path}"
        if self.used_ocr:
            s += " (OCR applied"
            s += f", +{self.ocr_supplement_pages} page(s) of OCR text appended)" if self.ocr_supplement_pages else ")"
        if self.fidelity:
            s += "\n" + self.fidelity.summary()
        return s

def _docx_to_pdf(docx_path: str, out_dir: str, timeout: int = 90) -> Optional[str]:
    """Convert DOCX to PDF using LibreOffice for fidelity checking."""
    if not CAPS["libreoffice"]:
        return None
    try:
        subprocess.run(
            ["soffice", "--headless", "--norestore", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
            check=True, timeout=timeout, capture_output=True,
        )
    except Exception as e:
        logger.warning(f"soffice conversion for verification failed: {e}")
        return None
    candidate = Path(out_dir) / (Path(docx_path).stem + ".pdf")
    return str(candidate) if candidate.exists() else None

def _compose_side_by_side(left_png: str, right_png: str, out_png: str, label_left="PDF", label_right="DOCX") -> str:
    """Create a side-by-side comparison image of two PNG files."""
    a, b = Image.open(left_png), Image.open(right_png)
    h = max(a.height, b.height)
    a = a.resize((int(a.width * h / a.height), h))
    b = b.resize((int(b.width * h / b.height), h))
    gap = 12
    composite = Image.new("RGB", (a.width + b.width + gap, h), "white")
    composite.paste(a, (0, 0))
    composite.paste(b, (a.width + gap, 0))
    composite.save(out_png)
    return out_png

def _autocrop(im: Image.Image, margin: int = 20) -> Image.Image:
    """Autocrop an image to remove surrounding whitespace."""
    from PIL import ImageChops
    bg = Image.new(im.mode, im.size, "white")
    diff = ImageChops.difference(im.convert("RGB"), bg.convert("RGB"))
    bbox = diff.getbbox()
    if not bbox:
        return im
    l, t, r, b = bbox
    l, t = max(l - margin, 0), max(t - margin, 0)
    r, b = min(r + margin, im.width), min(b + margin, im.height)
    return im.crop((l, t, r, b))

def _compose_grid(png_paths: List[str], labels: List[str], out_png: str, cols: int = 2) -> str:
    """Create a grid of images with labels."""
    from PIL import ImageDraw, ImageFont
    imgs = [_autocrop(Image.open(p)) for p in png_paths]
    tile_w = max(im.width for im in imgs)
    tile_h = max(im.height for im in imgs)
    label_h = 26
    rows = (len(imgs) + cols - 1) // cols
    grid = Image.new("RGB", (tile_w * cols, (tile_h + label_h) * rows), "white")
    draw = ImageDraw.Draw(grid)
    font = ImageFont.load_default()
    for idx, (im, label) in enumerate(zip(imgs, labels)):
        r, c = divmod(idx, cols)
        x, y = c * tile_w, r * (tile_h + label_h)
        draw.text((x + 6, y + 6), label, fill="black", font=font)
        grid.paste(im, (x, y + label_h))
    grid.save(out_png)
    return out_png

def debug_layout(pdf_path: str, page_index: int = 0, out_dir: Optional[str] = None,
                  dpi: int = 140) -> Optional[dict]:
    """Helps understand why a table or paragraph came out the way it did by showing the converter's internal detection stages."""
    out_dir = out_dir or tempfile.mkdtemp(prefix="debug_layout_")
    os.makedirs(out_dir, exist_ok=True)
    stem = Path(pdf_path).stem
    debug_pdf = str(Path(out_dir) / f"debug_{stem}.pdf")
    tmp_docx = str(Path(out_dir) / f"__debug_{stem}.docx")

    cv = _get_Converter()(pdf_path)
    try:
        cv.debug_page(page_index, docx_filename=tmp_docx, debug_pdf=debug_pdf,
                       layout_file=str(Path(out_dir) / "layout.json"))
    finally:
        cv.close()
    if not os.path.exists(debug_pdf):
        return None

    doc = fitz.open(debug_pdf)
    stages = {}
    for i, page in enumerate(doc):
        label = (page.get_text().strip().splitlines() or [f"stage_{i}"])[0]
        png_path = str(Path(out_dir) / f"debug_{stem}_{i}_{label.replace(' ', '_')}.png")
        page.get_pixmap(dpi=dpi).save(png_path)
        stages[label] = png_path
    doc.close()

    stages["composite"] = _compose_grid(list(stages.values()), list(stages.keys()),
                                         str(Path(out_dir) / f"debug_{stem}_grid.png"))
    return stages

def check_fidelity(pdf_path: str, docx_path: str, *, max_pages: int = 3, dpi: int = 120,
                    out_dir: Optional[str] = None, password: Optional[str] = None) -> FidelityReport:
    """Check the fidelity of a PDF to DOCX conversion by comparing visual similarity."""
    out_dir = out_dir or tempfile.mkdtemp(prefix="fidelity_")
    os.makedirs(out_dir, exist_ok=True)
    warnings: List[str] = []

    src = _open_pdf(pdf_path, password)
    pdf_pages = len(src)
    pdf_text = "".join(page.get_text() for page in src)

    page_images, comparisons = [], []
    docx_rendered_pages, docx_text = None, ""

    rendered_pdf = _docx_to_pdf(docx_path, out_dir)
    if rendered_pdf:
        out = fitz.open(rendered_pdf)
        docx_rendered_pages = len(out)
        docx_text = "".join(page.get_text() for page in out)
        for i in range(min(max_pages, pdf_pages, docx_rendered_pages)):
            pdf_png = str(Path(out_dir) / f"pdf_p{i+1}.png")
            docx_png = str(Path(out_dir) / f"docx_p{i+1}.png")
            src[i].get_pixmap(dpi=dpi).save(pdf_png)
            out[i].get_pixmap(dpi=dpi).save(docx_png)
            page_images.append((pdf_png, docx_png))
            comparisons.append(_compose_side_by_side(pdf_png, docx_png, str(Path(out_dir) / f"compare_p{i+1}.png")))
        out.close()
        if pdf_pages != docx_rendered_pages:
            warnings.append(f"Page count differs (PDF {pdf_pages} vs DOCX {docx_rendered_pages}).")
    else:
        warnings.append("LibreOffice ('soffice') not available — skipped visual comparison; "
                         "falling back to text extracted directly from the .docx.")
        try:
            d = DocxDocument(docx_path)
            parts = [p.text for p in d.paragraphs]
            for t in d.tables:
                for row in t.rows:
                    parts.extend(c.text for c in row.cells)
            docx_text = "\n".join(parts)
        except Exception as e:
            warnings.append(f"Could not read .docx text either: {e}")

    src.close()

    similarity = None
    if docx_text:
        a, b = pdf_text[:20000], docx_text[:20000]
        similarity = difflib.SequenceMatcher(None, a, b).ratio()
        if len(pdf_text.strip()) < SPARSE_CHAR_THRESHOLD:
            warnings.append("Source PDF has almost no extractable text (likely scanned), so the "
                             "text-similarity score isn't meaningful here — check the page images "
                             "instead, and see ocr_supplement_pages for recovered OCR text.")

    return FidelityReport(
        pdf_pages=pdf_pages, docx_rendered_pages=docx_rendered_pages, text_similarity=similarity,
        pdf_char_count=len(pdf_text), docx_char_count=len(docx_text),
        page_images=page_images, comparison_images=comparisons, warnings=warnings,
    )

def convert_pdf_to_docx(
    pdf_path: str, docx_path: str, *, start_page: int = 0, end_page: Optional[int] = None,
    pages: Optional[List[int]] = None, password: Optional[str] = None,
    force_ocr: bool = False, no_ocr: bool = False, append_ocr_text: bool = True,
    verify: bool = False, verify_max_pages: int = 3,
    **layout_kwargs,
) -> ConversionResult:
    """Convert a PDF file to DOCX format.

    Args:
        pdf_path: Path to input PDF file
        docx_path: Path to output DOCX file
        start_page: Start page (0-indexed)
        end_page: End page (exclusive)
        pages: Specific pages to convert (0-indexed list)
        password: Password for encrypted PDFs
        force_ocr: Force OCR on all pages
        no_ocr: Skip OCR even if needed
        append_ocr_text: Append OCR text as searchable supplement
        verify: Enable fidelity check
        verify_max_pages: Maximum pages to check for fidelity
        **layout_kwargs: Additional arguments passed to pdf2docx Converter

    Returns:
        ConversionResult object with conversion details
    """
    pdf_to_use, tmp_pdf, used_ocr = pdf_path, None, False

    try:
        needs_ocr = not no_ocr and (force_ocr or bool(sparse_page_indices(pdf_path, password=password)))
        if needs_ocr:
            tmp_fd, tmp_pdf = tempfile.mkstemp(suffix=".pdf")
            os.close(tmp_fd)
            if ocr_pdf(pdf_path, tmp_pdf, force_ocr=force_ocr):
                pdf_to_use, used_ocr = tmp_pdf, True
            else:
                logger.warning("OCR failed or unavailable; converting the original PDF as-is.")
                if os.path.exists(tmp_pdf):
                    os.unlink(tmp_pdf)
                tmp_pdf = None

        logger.info(f"Converting: {pdf_to_use} -> {docx_path}")
        cv = _get_Converter()(pdf_to_use, password=password)
        try:
            if pages is not None:
                cv.convert(docx_path, pages=pages, **layout_kwargs)
            else:
                cv.convert(docx_path, start=start_page, end=end_page, **layout_kwargs)
        finally:
            cv.close()
        logger.info(f"Successfully created: {docx_path}")

        supplement_pages = 0
        if used_ocr and append_ocr_text:
            candidates = sparse_page_indices(pdf_path, password=password)
            if pages is not None:
                candidates = [i for i in candidates if i in pages]
            else:
                candidates = [i for i in candidates if i >= start_page and (end_page is None or i < end_page)]
            supplement_pages = _append_ocr_supplement(docx_path, pdf_to_use, candidates)
            if supplement_pages:
                logger.info(f"Appended OCR text supplement for {supplement_pages} image-only page(s).")

        fidelity = check_fidelity(pdf_path, docx_path, max_pages=verify_max_pages, password=password) if verify else None
        return ConversionResult(pdf_path, docx_path, True, used_ocr=used_ocr,
                                 ocr_supplement_pages=supplement_pages, fidelity=fidelity)

    except Exception as e:
        logger.error(f"Conversion failed for {pdf_path}: {e}")
        return ConversionResult(pdf_path, None, False, used_ocr=used_ocr, error=str(e))
    finally:
        if tmp_pdf and os.path.exists(tmp_pdf):
            os.unlink(tmp_pdf)

def _append_ocr_supplement(docx_path: str, ocr_pdf_path: str, page_indices: List[int]) -> int:
    """Append OCR text as a searchable supplement to the DOCX file."""
    if not page_indices:
        return 0
    doc = fitz.open(ocr_pdf_path)
    texts = [(i, doc[i].get_text().strip()) for i in page_indices if 0 <= i < len(doc)]
    doc.close()
    texts = [(i, t) for i, t in texts if t]
    if not texts:
        return 0

    d = DocxDocument(docx_path)
    d.add_page_break()
    d.add_heading("OCR Text Supplement", level=1)
    note = d.add_paragraph(
        "The page(s) below were image-only in the source PDF. pdf2docx cannot "
        "reconstruct editable layout from a scan, so they appear above as a "
        "picture, so the text OCR recognized is reproduced here so it stays "
        "searchable and editable. Expect occasional recognition errors."
    )
    if note.runs:
        note.runs[0].italic = True
    for i, t in texts:
        d.add_heading(f"Page {i + 1} (from OCR)", level=2)
        d.add_paragraph(t)
    d.save(docx_path)
    return len(texts)

def convert_pdfs(
    input_path: str, output_dir: Optional[str] = None, *, start_page: int = 0,
    end_page: Optional[int] = None, pages: Optional[List[int]] = None,
    password: Optional[str] = None, force_ocr: bool = False, no_ocr: bool = False,
    verify: bool = False, verify_max_pages: int = 3, **layout_kwargs,
) -> List[ConversionResult]:
    """Convert PDF files in a directory to DOCX format.

    Args:
        input_path: Path to input PDF file or directory
        output_dir: Output directory for converted files (optional)
        start_page: Start page (0-indexed)
        end_page: End page (exclusive)
        pages: Specific pages to convert (0-indexed list)
        password: Password for encrypted PDFs
        force_ocr: Force OCR on all pages
        no_ocr: Skip OCR even if needed
        verify: Enable fidelity check
        verify_max_pages: Maximum pages to check for fidelity
        **layout_kwargs: Additional arguments passed to pdf2docx Converter

    Returns:
        List of ConversionResult objects
    """
    p = Path(input_path)
    if p.is_dir():
        pdf_paths = sorted(p.glob("*.pdf"))
    elif p.is_file() and p.suffix.lower() == ".pdf":
        pdf_paths = [p]
    else:
        raise ValueError(f"Invalid input: {input_path} — must be a PDF file or directory.")

    if not pdf_paths:
        logger.warning("No PDF files found.")
        return []

    out_dir = Path(output_dir) if output_dir else None
    if out_dir:
        out_dir.mkdir(parents=True, exist_ok=True)

    results = []
    for pdf_file in pdf_paths:
        docx_path = (out_dir / (pdf_file.stem + ".docx")) if out_dir else pdf_file.with_suffix(".docx")
        results.append(convert_pdf_to_docx(
            str(pdf_file), str(docx_path), start_page=start_page, end_page=end_page,
            pages=pages, password=password, force_ocr=force_ocr, no_ocr=no_ocr,
            verify=verify, verify_max_pages=verify_max_pages, **layout_kwargs,
        ))

    ok = sum(1 for r in results if r.success)
    logger.info(f"Done. {ok}/{len(pdf_paths)} file(s) converted.")
    return results